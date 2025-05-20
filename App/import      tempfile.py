import tempfile
import os
import time
import json
import re
import subprocess
from io import BytesIO
from datetime import datetime
import traceback

from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.models import User, auth
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse, StreamingHttpResponse
from django.core.files.storage import FileSystemStorage
import datetime

# Import models
from .models import Profile, Subtitle
from django.db.models import Sum

# Third-party libraries
import google.generativeai as genai
import torch
from moviepy import VideoFileClip
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

 
genai.configure(api_key='AIzaSyDSgnW6zyeUTEcGqyJJcKEhOJZXOHzz6uQ')

# Language mapping for translation and detection
language_names = {
    'auto': 'Auto Detect',
    'af': 'Afrikaans',
    'sq': 'Albanian',
    'am': 'Amharic',
    'ar': 'Arabic',
    'hy': 'Armenian',
    'az': 'Azerbaijani',
    'eu': 'Basque',
    'be': 'Belarusian',
    'bn': 'Bengali',
    'bs': 'Bosnian',
    'bg': 'Bulgarian',
    'ca': 'Catalan',
    'ceb': 'Cebuano',
    'ny': 'Chichewa',
    'zh': 'Chinese',
    'hr': 'Croatian',
    'cs': 'Czech',
    'da': 'Danish',
    'nl': 'Dutch',
    'en': 'English',
    'eo': 'Esperanto',
    'et': 'Estonian',
    'fi': 'Finnish',
    'fr': 'French',
    'gl': 'Galician',
    'ka': 'Georgian',
    'de': 'German',
    'el': 'Greek',
    'gu': 'Gujarati',
    'ht': 'Haitian Creole',
    'ha': 'Hausa',
    'haw': 'Hawaiian',
    'he': 'Hebrew',
    'hi': 'Hindi',
    'hmn': 'Hmong',
    'hu': 'Hungarian',
    'is': 'Icelandic',
    'ig': 'Igbo',
    'id': 'Indonesian',
    'ga': 'Irish',
    'it': 'Italian',
    'ja': 'Japanese',
    'jv': 'Javanese',
    'kn': 'Kannada',
    'kk': 'Kazakh',
    'km': 'Khmer',
    'ko': 'Korean',
    'ku': 'Kurdish',
    'ky': 'Kyrgyz',
    'lo': 'Lao',
    'la': 'Latin',
    'lv': 'Latvian',
    'lt': 'Lithuanian',
    'lb': 'Luxembourgish',
    'mk': 'Macedonian',
    'mg': 'Malagasy',
    'ms': 'Malay',
    'ml': 'Malayalam',
    'mt': 'Maltese',
    'mi': 'Maori',
    'mr': 'Marathi',
    'mn': 'Mongolian',
    'my': 'Myanmar (Burmese)',
    'ne': 'Nepali',
    'no': 'Norwegian',
    'ps': 'Pashto',
    'fa': 'Persian',
    'pl': 'Polish',
    'pt': 'Portuguese',
    'pa': 'Punjabi',
    'ro': 'Romanian',
    'ru': 'Russian',
    'sm': 'Samoan',
    'gd': 'Scots Gaelic',
    'sr': 'Serbian',
    'st': 'Sesotho',
    'sn': 'Shona',
    'sd': 'Sindhi',
    'si': 'Sinhala',
    'sk': 'Slovak',
    'sl': 'Slovenian',
    'so': 'Somali',
    'es': 'Spanish',
    'su': 'Sundanese',
    'sw': 'Swahili',
    'sv': 'Swedish',
    'tg': 'Tajik',
    'ta': 'Tamil',
    'te': 'Telugu',
    'th': 'Thai',
    'tr': 'Turkish',
    'uk': 'Ukrainian',
    'ur': 'Urdu',
    'uz': 'Uzbek',
    'vi': 'Vietnamese',
    'cy': 'Welsh',
    'xh': 'Xhosa',
    'yi': 'Yiddish',
    'yo': 'Yoruba',
    'zu': 'Zulu'
}

# Create reverse mapping of language names to codes
language_names_reverse = {v.lower(): k for k, v in language_names.items()}

# Helper Functions
def ms_to_srt_time(ms):
    """
    Convert milliseconds to SRT time format (HH:MM:SS,mmm)
    Fixed to ensure accurate timestamp conversion
    """
    # Ensure ms is a number
    try:
        ms = float(ms)
    except (ValueError, TypeError):
        ms = 0
    
    # Prevent negative values
    ms = max(0, ms)
    
    seconds, ms = divmod(ms, 1000)
    minutes, seconds = divmod(seconds, 60)
    hours, minutes = divmod(minutes, 60)
    
    # Ensure all values are integers for formatting
    return f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d},{int(ms):03d}"

def srt_time_to_ms(time_str):
    """
    Convert SRT time format (HH:MM:SS,mmm) to milliseconds
    """
    try:
        # Handle both comma and period as decimal separators
        time_str = time_str.replace('.', ',')
        
        # Split into hours, minutes, seconds, and milliseconds
        hours, minutes, seconds_ms = time_str.split(':')
        seconds, ms = seconds_ms.split(',')
        
        # Convert to milliseconds
        total_ms = (int(hours) * 3600 + int(minutes) * 60 + int(seconds)) * 1000 + int(ms)
        return total_ms
    except Exception as e:
        print(f"Error converting time to ms: {e}")
        return 0

def apply_time_offset(segments, offset_ms):
    """
    Apply a time offset (in milliseconds) to all subtitle segments
    Positive offset delays subtitles, negative offset makes them appear earlier
    """
    adjusted_segments = []
    
    for segment in segments:
        # Create a new segment with adjusted timestamps
        adjusted_segment = segment.copy()
        
        # Apply offset
        start_ms = segment['start'] + offset_ms
        end_ms = segment['end'] + offset_ms
        
        # Ensure timestamps don't go negative
        adjusted_segment['start'] = max(0, start_ms)
        adjusted_segment['end'] = max(adjusted_segment['start'] + 500, end_ms)  # Ensure at least 500ms duration
        
        adjusted_segments.append(adjusted_segment)
    
    return adjusted_segments

def reduce_audio_noise(input_path, output_path):
    """
    Use FFmpeg to reduce noise from audio file
    
    Parameters:
    - input_path: Path to the input audio file
    - output_path: Path to save the noise-reduced audio file
    
    Returns:
    - True if successful, False otherwise
    """
    try:
        # FFmpeg command for noise reduction
        # Using highpass filter to remove low frequency noise (below 200Hz)
        # Using lowpass filter to remove high frequency noise (above 3000Hz)
        # Using compand filter to normalize volume and reduce background noise
        ffmpeg_cmd = [
            'ffmpeg', '-y', '-i', input_path,
            '-af', 'highpass=f=200,lowpass=f=3000,compand=0.3|0.3:1|1:-90/-60|-60/-40|-40/-30|-20/-20:6:0:-90:0.2',
            '-ar', '44100',  # Set sample rate
            output_path
        ]
        
        # Execute the command
        process = subprocess.Popen(ffmpeg_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()
        
        if process.returncode != 0:
            print(f"FFmpeg error: {stderr.decode()}")
            return False
            
        return True
    except Exception as e:
        print(f"Error reducing audio noise: {e}")
        return False

def sync_subtitles_with_audio(audio_path, srt_path, output_path):
    """
    Use FFsubsync to synchronize subtitles with audio
    
    Parameters:
    - audio_path: Path to the audio file
    - srt_path: Path to the input SRT file
    - output_path: Path to save the synchronized SRT file
    
    Returns:
    - True if successful, False otherwise
    """
    try:
        # FFsubsync command to sync subtitles with audio
        ffsubsync_cmd = [
            'ffs', audio_path, '-i', srt_path, '-o', output_path
        ]
        
        # Execute the command
        process = subprocess.Popen(ffsubsync_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()
        
        if process.returncode != 0:
            print(f"FFsubsync error: {stderr.decode()}")
            return False
            
        return True
    except Exception as e:
        print(f"Error syncing subtitles: {e}")
        return False

def detect_language_with_gemini(audio_path):
    """Detect language using Gemini API"""
    try:
        # Load audio file into memory
        with open(audio_path, 'rb') as audio_file:
            audio_data = audio_file.read()
        
        # Create a Gemini model for audio processing
        model = genai.GenerativeModel('gemini-2.5-flash-preview-04-17')
        
        # Prepare prompt for language detection
        prompt = "Please identify the language spoken in this audio file. Respond with only the language name."
        
        # Make API request with audio content
        response = model.generate_content([prompt, {"mime_type": "audio/mp3", "data": audio_data}])
        
        # Process response to get language name
        detected_language = response.text.strip().lower()
        
        # Map language name to code
        detected_code = language_names_reverse.get(detected_language, 'en')  # Default to English if not found
        full_name = language_names.get(detected_code, detected_language)
        
        print(f"Detected language: {full_name} ({detected_code})")
        return (detected_code, full_name)
    except Exception as e:
        print(f"Language detection exception: {str(e)}")
        traceback.print_exc()
        return ('en', 'English')  # Default to English on error

def transcribe_with_gemini(audio_path, language=None, task="transcribe", target_language=None):
    """
    Transcribe or translate audio using Gemini API
    
    Parameters:
    - audio_path: Path to the audio file
    - language: The source language code (optional)
    - task: "transcribe" or "translate"
    - target_language: Target language for translation
    
    Returns:
    - A list of dictionaries with start_time, end_time, and text
    """
    try:
        # Load audio file into memory
        with open(audio_path, 'rb') as audio_file:
            audio_data = audio_file.read()
        
        # Create a Gemini model for audio processing
        model = genai.GenerativeModel('gemini-2.5-flash-preview-04-17')
        
        # Prepare prompt based on task
        if task == "translate":
            target_lang_name = language_names.get(target_language, "English")
            prompt = f"""Please transcribe the following audio and then translate it to {target_lang_name}. 
            Return the result as a JSON array where each element has 'start' (timestamp in milliseconds), 
            'end' (timestamp in milliseconds), and 'text' (the translated text).
            CRITICAL: The timestamps MUST be accurate and synchronized with the audio content.
            Each segment should be between 1-5 seconds long for proper subtitle display.
            Start timestamps at 0 and ensure they match when the words are actually spoken.
            DO NOT start timestamps before the actual speech begins."""
        else:
            source_lang_name = language_names.get(language, "the detected language") if language != 'auto' else "the detected language"
            prompt = f"""Please transcribe the following audio in {source_lang_name}. 
            Return the result as a JSON array where each element has 'start' (timestamp in milliseconds), 
            'end' (timestamp in milliseconds), and 'text' (the transcribed text).
            CRITICAL: The timestamps MUST be accurate and synchronized with the audio content.
            Each segment should be between 1-5 seconds long for proper subtitle display.
            Start timestamps at 0 and ensure they match when the words are actually spoken.
            DO NOT start timestamps before the actual speech begins."""
        
        print(f"Using prompt: {prompt}")
        
        # Make API request with audio content
        response = model.generate_content([prompt, {"mime_type": "audio/mp3", "data": audio_data}])
        
        # Extract JSON from response
        result_text = response.text.strip()
        # Handle case where response might be wrapped in code blocks
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', result_text)
        if json_match:
            json_str = json_match.group(1)
        else:
            json_str = result_text
            
        # Parse JSON result
        try:
            transcribed_segments = json.loads(json_str)
            
            # Ensure proper format - convert string timestamps to integers if needed
            for segment in transcribed_segments:
                if isinstance(segment.get('start'), str):
                    # Handle potential floating point strings
                    segment['start'] = int(float(segment['start']) * 1000)
                if isinstance(segment.get('end'), str):
                    segment['end'] = int(float(segment['end']) * 1000)
                
                # Validate timestamps
                if segment['start'] >= segment['end']:
                    # Fix invalid timestamps by adding a small duration
                    segment['end'] = segment['start'] + 2000  # Add 2 seconds
                
                # Ensure text exists
                if 'text' not in segment or not segment['text']:
                    segment['text'] = "[No text]"
                    
            # Sort segments by start time
            transcribed_segments.sort(key=lambda x: x['start'])
            
            # Check for overlapping segments and fix them
            for i in range(1, len(transcribed_segments)):
                if transcribed_segments[i]['start'] < transcribed_segments[i-1]['end']:
                    transcribed_segments[i]['start'] = transcribed_segments[i-1]['end'] + 100  # Add 100ms gap
                    
                    # If this causes start to be after end, adjust end too
                    if transcribed_segments[i]['start'] >= transcribed_segments[i]['end']:
                        transcribed_segments[i]['end'] = transcribed_segments[i]['start'] + 2000  # Add 2 seconds
            
            # Apply a default offset to ensure subtitles are properly synchronized
            # This is a common fix for subtitle timing issues - adjust as needed
            default_offset_ms = 500  # 0.5 second delay - reduced from original since we'll use ffsubsync later
            transcribed_segments = apply_time_offset(transcribed_segments, default_offset_ms)
            
            return transcribed_segments
        except json.JSONDecodeError:
            print(f"Failed to parse JSON response: {result_text[:500]}...")
            # Fall back to creating a single segment with the entire text
            return [{
                'start': 500,  # Start after 0.5 seconds to avoid early display
                'end': 10500,  # 10 seconds default
                'text': result_text.strip()
            }]
            
    except Exception as e:
        print(f"Transcription exception: {str(e)}")
        traceback.print_exc()
        return []

def convert_srt_to_plain_text(srt_content):
    """Convert SRT format to plain text"""
    plain_text = ''
    lines = srt_content.split('\n')
    current_text = ''

    for i in range(len(lines)):
        line = lines[i].strip()

        # Skip index numbers and timestamp lines
        if re.match(r'^\d+$', line) or '-->' in line:
            continue

        # Add text content
        if line:
            current_text += line + ' '
        elif current_text:
            # End of a subtitle block
            plain_text += current_text.strip() + '\n\n'
            current_text = ''

    # Add the last block if exists
    if current_text:
        plain_text += current_text.strip()

    return plain_text

def get_or_create_user_profile(user):
    """Get or create a user profile"""
    try:
        return Profile.objects.get(user=user)
    except Profile.DoesNotExist:
        profile = Profile.objects.create(user=user)
        profile.save()
        return profile

# API endpoints for React/Next.js integration
def api_get_subtitles(request, subtitle_id):
    """API endpoint to get a subtitle file for React frontend"""
    try:
        subtitle = Subtitle.objects.get(id=subtitle_id)
        fs = FileSystemStorage()
        
        # Prepare response data
        subtitle_url = fs.url(subtitle.subtitle.name) if subtitle.subtitle else None
        vtt_url = fs.url(subtitle.webVit.name) if subtitle.webVit else None
        audio_url = fs.url(subtitle.audio.name) if subtitle.audio else None
        
        return JsonResponse({
            'id': subtitle.id,
            'audio_url': audio_url,
            'subtitle_url': subtitle_url, 
            'vtt_url': vtt_url,
            'source_lang': subtitle.source_lang,
            'target_lang': subtitle.target_lang,
            'created_at': subtitle.created_at.isoformat(),
            'updated_at': subtitle.updated_at.isoformat() if subtitle.updated_at else None,
            'status': subtitle.status
        })
        
    except Subtitle.DoesNotExist:
        return JsonResponse({'error': 'Subtitle not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def api_list_subtitles(request):
    """API endpoint to list user's subtitles for React frontend"""
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Authentication required'}, status=401)
        
    try:
        user_profile = get_or_create_user_profile(request.user)
        subtitles = Subtitle.objects.filter(user_profile=user_profile, is_deleted=False)
        
        subtitle_list = []
        for sub in subtitles:
            subtitle_list.append({
                'id': sub.id,
                'filename': os.path.basename(sub.subtitle.name) if sub.subtitle else None,
                'source_lang': sub.source_lang,
                'target_lang': sub.target_lang,
                'created_at': sub.created_at.isoformat(),
                'status': sub.status
            })
            
        return JsonResponse({'subtitles': subtitle_list})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def api_adjust_subtitle(request, subtitle_id):
    """API endpoint to adjust subtitle timing for React frontend"""
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Authentication required'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
        
    try:
        data = json.loads(request.body)
        offset = data.get('offset', 0)  # Get offset in seconds
        
        # Convert to milliseconds
        try:
            offset_ms = int(float(offset) * 1000)
        except ValueError:
            return JsonResponse({'error': 'Invalid offset value'}, status=400)
            
        # Get the subtitle
        subtitle = Subtitle.objects.get(id=subtitle_id)
        
        # Check if user has permission
        if subtitle.user_profile.user != request.user:
            return JsonResponse({'error': 'Permission denied'}, status=403)
            
        fs = FileSystemStorage()
        subtitle_path = subtitle.subtitle.path if subtitle.subtitle else None
        webvit_path = subtitle.webVit.path if subtitle.webVit else None
        
        if not subtitle_path or not webvit_path:
            return JsonResponse({'error': 'Subtitle files not found'}, status=404)
            
        # Read the SRT file
        with open(subtitle_path, 'r', encoding='utf-8') as f:
            srt_content = f.read()
            
        # Apply offset to timestamps
        lines = srt_content.split('\n')
        adjusted_lines = []
        
        for line in lines:
            line = line.strip()
            adjusted_lines.append(line)
            
            # Check if this is a timestamp line
            if '-->' in line:
                start_time, end_time = line.split(' --> ')
                
                # Apply offset
                start_ms = srt_time_to_ms(start_time) + offset_ms
                end_ms = srt_time_to_ms(end_time) + offset_ms
                
                # Ensure times don't go negative
                start_ms = max(0, start_ms)
                end_ms = max(start_ms + 500, end_ms)
                
                # Format back to SRT time
                adjusted_start = ms_to_srt_time(start_ms)
                adjusted_end = ms_to_srt_time(end_ms)
                
                # Replace with adjusted timestamps
                adjusted_lines[-1] = f"{adjusted_start} --> {adjusted_end}"
                
        # Join lines back together
        adjusted_srt = '\n'.join(adjusted_lines)
        
        # Write the adjusted SRT file
        with open(subtitle_path, 'w', encoding='utf-8') as f:
            f.write(adjusted_srt)
            
        # Convert SRT to WebVTT for the webVit file
        vtt_content = "WEBVTT\n\n" + adjusted_srt.replace(',', '.')
        with open(webvit_path, 'w', encoding='utf-8') as f:
            f.write(vtt_content)
            
        # Update the last modified time
        subtitle.updated_at = datetime.now()
        subtitle.save()
        
        # Prepare URLs for response
        subtitle_url = fs.url(subtitle.subtitle.name)
        vtt_url = fs.url(subtitle.webVit.name)
        
        return JsonResponse({
            'status': 'success',
            'message': f'Subtitle timing adjusted by {offset} seconds',
            'subtitle_url': subtitle_url,
            'vtt_url': vtt_url
        })
        
    except Subtitle.DoesNotExist:
        return JsonResponse({'error': 'Subtitle not found'}, status=404)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# View Functions
def index(request):
    """Home page view"""
    return render(request, 'index.html')

def login(request):
    """User login view"""
    if request.method == 'POST':
        email = request.POST['email'].strip()
        password = request.POST['password'].strip()
        
        # Authenticate user
        user = auth.authenticate(username=email, password=password)
        
        if user is not None:
            auth.login(request, user)
            messages.success(request, f"Successfully logged in as {user}")
            return redirect('upload')
        else:
            messages.error(request, 'Invalid Credentials')
            return redirect('login')
        
    return render(request, 'login.html')

def register(request):
    """User registration view"""
    if request.method == 'POST':
        username = request.POST['name'].strip()
        email = request.POST['email'].strip()
        password = request.POST['password'].strip()
        password2 = request.POST['confirm_password'].strip()
        
        if password == password2:
            if User.objects.filter(email=email).exists():
                messages.error(request, 'Email Already Exists')
            elif User.objects.filter(username=username).exists():
                messages.error(request, 'Username Already Exists')
            else:
                # Create user using Django's create_user method
                user = User.objects.create_user(
                    username=username,
                    email=email,
                    password=password
                )
                user.save()
                
                # Create profile for the new user
                profile = Profile.objects.create(user=user)
                profile.save()
                
                messages.success(request, 'Successfully Registered')
                return redirect('login')
        else:
            messages.error(request, 'Password does not match')

    return render(request, 'register.html')

def logout(request):
    """User logout view"""
    auth.logout(request)
    return redirect('login')

@login_required(login_url='/login')
def dashboard(request):
    """User dashboard view showing all transcriptions"""
    user_profile = get_or_create_user_profile(request.user)
    total_transcriptions = Subtitle.objects.filter(user_profile=user_profile, is_deleted=False)
    transcriptions_length = total_transcriptions.count()
    
    total_seconds = total_transcriptions.aggregate(total=Sum('duration'))['total'] or 0
    total_duration = str(datetime.timedelta(seconds=int(total_seconds)))
    just_seconds = int(total_seconds)
    
    
    pending = total_transcriptions.filter(status='pending').count()
    processing = total_transcriptions.filter(status='processing').count()
    completed = total_transcriptions.filter(status='completed').count()
    failed = total_transcriptions.filter(status='failed').count()
    
    context = {
        'total_transcriptions': total_transcriptions,
        'transcriptions_length': transcriptions_length,
        'pending': pending,
        'processing': processing,
        'completed': completed,
        'failed': failed,
        'just_seconds':just_seconds
    }
    
    return render(request, 'dashboard.html', context)

@login_required(login_url='/login')    
def upload(request):
    """Handle video upload, transcription and translation"""
    user_profile = get_or_create_user_profile(request.user)
    
    # Get available languages for the form
    languages = sorted([(code, name) for code, name in language_names.items() 
                        if code != 'auto'], key=lambda x: x[1])
    
    if request.method == "POST":
        video = request.FILES.get('video')
        if not video:
            messages.error(request, 'Please upload a video file')
            return render(request, 'upload.html', {'languages': languages})
            
        source_lang = request.POST.get('source-language', 'auto')
        target_lang = request.POST.get('target-language', 'en')
        output_format = request.POST.get('format', 'srt')
        include_timestamps = request.POST.get('timestamps') == 'on'
        
        # Get subtitle offset (if provided)
        subtitle_offset = request.POST.get('subtitle_offset', '0')  # Default to 0 seconds since we'll use ffsubsync
        try:
            subtitle_offset_ms = int(float(subtitle_offset) * 1000)  # Convert to milliseconds
        except ValueError:
            subtitle_offset_ms = 0  # Default to 0 seconds if invalid
        
        processing_start_time = time.time()
        fs = FileSystemStorage()
        
        # Create a subtitle entry with pending status
        subtitle = Subtitle.objects.create(
            user_profile=user_profile,
            source_lang='Auto Detect' if source_lang == 'auto' else language_names.get(source_lang, source_lang),
            target_lang=language_names.get(target_lang, target_lang),
            status='processing'
        )
        
        try:
            # Process video in a temporary file
            with tempfile.NamedTemporaryFile(suffix=os.path.splitext(video.name)[1], delete=False) as temp_video:
                # Save uploaded video to temp file
                for chunk in video.chunks():
                    temp_video.write(chunk)
                
                # Close the file so it can be accessed by moviepy
                temp_video.close()
                
                video_temp_path = temp_video.name
                base_name = os.path.splitext(os.path.basename(video.name))[0]
                
                # Extract audio and save it temporarily for processing
                temp_audio_path = os.path.join(tempfile.gettempdir(), f"{base_name}_temp_{int(time.time())}.mp3")
                
                # Load video and extract audio
                clip = VideoFileClip(video_temp_path)
                clip.audio.write_audiofile(temp_audio_path)
                audio_duration = clip.duration  # Get duration in seconds
                # Free up resources
                clip.close()
                
                # Create final audio output path
                audio_filename = f"{base_name}_{int(time.time())}.mp3"
                audio_path = os.path.join(fs.location, audio_filename)
                
                # Apply noise reduction with FFmpeg
                print("Applying noise reduction with FFmpeg...")
                noise_reduction_success = reduce_audio_noise(temp_audio_path, audio_path)
                
                if not noise_reduction_success:
                    print("Noise reduction failed, using original audio file")
                    # If noise reduction fails, copy the original audio file
                    with open(temp_audio_path, 'rb') as src, open(audio_path, 'wb') as dst:
                        dst.write(src.read())
                
                # Update subtitle with audio file
                subtitle.audio = audio_filename
                subtitle.duration = f"{audio_duration:.2f} seconds"
                subtitle.save()
                
                # Language detection
                detected_lang_code = None
                detected_lang_name = None
                
                # Auto-detect source language if needed
                if source_lang == 'auto':
                    detected_lang_code, detected_lang_name = detect_language_with_gemini(audio_path)
                    source_lang = detected_lang_code  # Use detected language for further processing
                
                # Get full language names
                source_language_name = detected_lang_name if source_lang == 'auto' else language_names.get(source_lang, source_lang)
                target_language_name = language_names.get(target_lang, target_lang)
                
                # Update subtitle with detected language
                if source_lang == 'auto' and detected_lang_name:
                    subtitle.source_lang = f"{detected_lang_name} (Auto Detected)"
                else:
                    subtitle.source_lang = source_language_name
                subtitle.save()
                
                # Determine if we need to translate or just transcribe
                task = "transcribe"
                if source_lang != target_lang:
                    task = "translate"
                    print(f"Using Gemini to translate from {source_lang} to {target_lang}")
                else:
                    print(f"Using Gemini to transcribe in {source_lang}")

                # Transcribe the audio
                transcribed_segments = transcribe_with_gemini(audio_path, source_lang, task, target_lang)
                
                # Create temporary SRT file from Gemini transcription
                temp_srt_path = os.path.join(tempfile.gettempdir(), f"{base_name}_temp_{int(time.time())}.srt")
                with open(temp_srt_path, 'w', encoding='utf-8') as f:
                    for i, sub in enumerate(transcribed_segments, start=1):
                        f.write(f"{i}\n")
                        f.write(f"{ms_to_srt_time(sub['start'])} --> {ms_to_srt_time(sub['end'])}\n")
                        f.write(f"{sub['text']}\n\n")
                
                # Use FFsubsync to improve subtitle synchronization
                print("Syncing subtitles with audio using FFsubsync...")
                final_srt_path = os.path.join(tempfile.gettempdir(), f"{base_name}_final_{int(time.time())}.srt")
                sync_success = sync_subtitles_with_audio(audio_path, temp_srt_path, final_srt_path)
                
                # If ffsubsync fails, use the original timing with manual offset
                if not sync_success:
                    print("FFsubsync failed, using original timing with manual offset")
                    final_srt_path = temp_srt_path
                    
                    # Apply additional user-specified offset if provided
                    if subtitle_offset_ms != 0:
                        print(f"Applying manual offset of {subtitle_offset_ms}ms")
                        # Read the temporary SRT file
                        with open(temp_srt_path, 'r', encoding='utf-8') as f:
                            srt_content = f.read()
                            
                        # Apply offset to timestamps
                        lines = srt_content.split('\n')
                        adjusted_lines = []
                        
                        for line in lines:
                            line = line.strip()
                            adjusted_lines.append(line)
                            
                            # Check if this is a timestamp line
                            if '-->' in line:
                                start_time, end_time = line.split(' --> ')
                                
                                # Apply offset
                                start_ms = srt_time_to_ms(start_time) + subtitle_offset_ms
                                end_ms = srt_time_to_ms(end_time) + subtitle_offset_ms
                                
                                # Ensure times don't go negative
                                start_ms = max(0, start_ms)
                                end_ms = max(start_ms + 500, end_ms)
                                
                                # Format back to SRT time
                                adjusted_start = ms_to_srt_time(start_ms)
                                adjusted_end = ms_to_srt_time(end_ms)
                                
                                # Replace with adjusted timestamps
                                adjusted_lines[-1] = f"{adjusted_start} --> {adjusted_end}"
                                
                        # Join lines back together
                        adjusted_srt = '\n'.join(adjusted_lines)
                        
                        # Write the adjusted SRT file
                        with open(final_srt_path, 'w', encoding='utf-8') as f:
                            f.write(adjusted_srt)
                
                # Read the final synchronized SRT content
                with open(final_srt_path, 'r', encoding='utf-8') as f:
                    synced_srt_content = f.read()
                
                # Generate output based on selected format
                output_filename = f"{base_name}_{target_lang}_{int(time.time())}"
                output_path = None
                
                if output_format == 'txt':
                    # Plain text output
                    output_filename += '.txt'
                    output_path = os.path.join(fs.location, output_filename)
                    
                    # Convert SRT to plain text
                    plain_text = convert_srt_to_plain_text(synced_srt_content)
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        if include_timestamps:
                            # Parse SRT and keep timestamps
                            lines = synced_srt_content.split('\n')
                            current_timestamp = ""
                            for i in range(len(lines)):
                                line = lines[i].strip()
                                
                                if '-->' in line:
                                    start_time, end_time = line.split(' --> ')
                                    current_timestamp = f"[{start_time.replace(',', '.')} - {end_time.replace(',', '.')}] "
                                elif line and not line.isdigit() and i > 0 and not '-->' in lines[i-1]:
                                    f.write(f"{current_timestamp}{line}\n\n")
                                    current_timestamp = ""
                        else:
                            f.write(plain_text)
                
                elif output_format == 'srt':
                    # SRT subtitle format
                    output_filename += '.srt'
                    output_path = os.path.join(fs.location, output_filename)
                    
                    # Copy the final SRT file
                    with open(final_srt_path, 'r', encoding='utf-8') as src:
                        with open(output_path, 'w', encoding='utf-8') as dst:
                            dst.write(src.read())
                
                elif output_format == 'vtt':
                    # VTT subtitle format
                    output_filename += '.vtt'
                    output_path = os.path.join(fs.location, output_filename)
                    
                    # Convert SRT to VTT
                    vtt_content = "WEBVTT\n\n" + synced_srt_content.replace(',', '.')
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(vtt_content)
                
                elif output_format == 'docx':
                    # Word document format
                    try:
                        output_filename += '.docx'
                        output_path = os.path.join(fs.location, output_filename)
                        
                        doc = Document()
                        doc.add_heading(f"Transcript: {base_name}", 0)
                        
                        # Parse SRT content
                        lines = synced_srt_content.split('\n')
                        current_timestamp = ""
                        current_text = ""
                        
                        for i in range(len(lines)):
                            line = lines[i].strip()
                            
                            if '-->' in line:
                                start_time, end_time = line.split(' --> ')
                                current_timestamp = f"[{start_time.replace(',', '.')} - {end_time.replace(',', '.')}] "
                            elif line and not line.isdigit() and i > 0 and not '-->' in lines[i-1]:
                                paragraph = doc.add_paragraph()
                                if include_timestamps:
                                    paragraph.add_run(current_timestamp).bold = True
                                paragraph.add_run(line)
                                current_timestamp = ""
                        
                        doc.save(output_path)
                        
                    except ImportError:
                        # Fallback to text if python-docx not installed
                        output_filename = f"{base_name}_{target_lang}_{int(time.time())}.txt"
                        output_path = os.path.join(fs.location, output_filename)
                        
                        # Convert SRT to plain text
                        plain_text = convert_srt_to_plain_text(synced_srt_content)
                        
                        with open(output_path, 'w', encoding='utf-8') as f:
                            f.write(f"Transcript: {base_name}\n\n")
                            f.write(plain_text)
                
                elif output_format == 'pdf':
                    # PDF format
                    try:
                        output_filename += '.pdf'
                        output_path = os.path.join(fs.location, output_filename)
                        
                        doc = canvas.Canvas(output_path, pagesize=letter)
                        width, height = letter
                        
                        # Add title
                        doc.setFont("Helvetica-Bold", 16)
                        doc.drawString(72, height - 72, f"Transcript: {base_name}")
                        
                        # Add content
                        doc.setFont("Helvetica", 12)
                        y_position = height - 100
                        
                        # Parse SRT content
                        lines = synced_srt_content.split('\n')
                        current_timestamp = ""
                        current_text = ""
                        
                        for i in range(len(lines)):
                            line = lines[i].strip()
                            
                            if '-->' in line:
                                if include_timestamps:
                                    start_time, end_time = line.split(' --> ')
                                    current_timestamp = f"[{start_time.replace(',', '.')} - {end_time.replace(',', '.')}] "
                                else:
                                    current_timestamp = ""
                            elif line and not line.isdigit() and i > 0 and not '-->' in lines[i-1]:
                                text_line = current_timestamp + line
                                
                                # Simple text wrapping
                                words = text_line.split()
                                line_text = ""
                                for word in words:
                                    test_line = line_text + " " + word if line_text else word
                                    if doc.stringWidth(test_line, "Helvetica", 12) < width - 144:
                                        line_text = test_line
                                    else:
                                        doc.drawString(72, y_position, line_text)
                                        y_position -= 15
                                        line_text = word
                                        
                                        # New page if needed
                                        if y_position < 72:
                                            doc.showPage()
                                            doc.setFont("Helvetica", 12)
                                            y_position = height - 72
                                
                                if line_text:
                                    doc.drawString(72, y_position, line_text)
                                    y_position -= 15
                                
                                # Add spacing between subtitles
                                y_position -= 10
                                
                                # New page if needed
                                if y_position < 72:
                                    doc.showPage()
                                    doc.setFont("Helvetica", 12)
                                    y_position = height - 72
                                    
                                current_timestamp = ""
                        
                        doc.save()
                        
                    except Exception as e:
                        print(f"PDF generation error: {e}")
                        # Fallback to text
                        output_filename = f"{base_name}_{target_lang}_{int(time.time())}.txt"
                        output_path = os.path.join(fs.location, output_filename)
                        
                        # Convert SRT to plain text
                        plain_text = convert_srt_to_plain_text(synced_srt_content)
                        
                        with open(output_path, 'w', encoding='utf-8') as f:
                            f.write(f"Transcript: {base_name}\n\n")
                            f.write(plain_text)
                
                # For browser viewing, always create a VTT file
                vtt_filename = f"{base_name}_{target_lang}_{int(time.time())}.vtt"
                vtt_path = os.path.join(fs.location, vtt_filename)
                
                # Convert SRT to VTT
                vtt_content = "WEBVTT\n\n" + synced_srt_content.replace(',', '.')
                with open(vtt_path, 'w', encoding='utf-8') as f:
                    f.write(vtt_content)
                
                # Clean up temporary files
                try:
                    os.unlink(video_temp_path)
                    os.unlink(temp_audio_path)
                    os.unlink(temp_srt_path)
                    os.unlink(final_srt_path)
                except Exception as e:
                    print(f"Error removing temp files: {e}")
                
                # Update subtitle record
                processing_duration = time.time() - processing_start_time
                
                subtitle.subtitle = output_filename
                subtitle.webVit = vtt_filename
                subtitle.actions = 'transcribed video with noise reduction and synchronized subtitles'
                subtitle.status = 'completed'
                subtitle.is_completed = True
                subtitle.processing_duration = f"{processing_duration:.2f} seconds"
                subtitle.save()
                
                # Context dictionary
                context = {
                    'user_profile': user_profile,
                    'user_subtitle': subtitle,
                    'audio_url': fs.url(audio_filename),
                    'subtitle_url': fs.url(output_filename),
                    'webVit_url': fs.url(vtt_filename),
                    'output_format': output_format,
                    'success': True,
                    'source_lang': source_lang,
                    'target_lang': target_lang,
                    'source_language_name': source_language_name,
                    'target_language_name': target_language_name,
                    'processing_duration': processing_duration,
                    'duration': f"{audio_duration:.2f} seconds",
                    'languages': languages
                }
                
                return render(request, 'upload.html', context)
                
        except Exception as e:
            # Handle errors
            subtitle.status = 'failed'
            subtitle.save()
            
            print(f"Error processing video: {str(e)}")
            traceback.print_exc()
            messages.error(request, f"Error processing video: {str(e)}")
            
            return render(request, 'upload.html', {'languages': languages})
    
    # GET request - show upload form
    return render(request, 'upload.html', {'languages': languages})

@login_required(login_url='/login')  
def view_subtitle(request, pk):
    """View a specific subtitle"""
    try:
        user_subtitle = Subtitle.objects.get(id=pk)
        fs = FileSystemStorage()
        
        # Check if user has permission to view this subtitle
        if user_subtitle.user_profile.user != request.user:
            messages.error(request, "You don't have permission to view this subtitle")
            return redirect('dashboard')
        
        subtitle_url = fs.url(user_subtitle.subtitle.name) if user_subtitle.subtitle else None
        audio_url = fs.url(user_subtitle.audio.name) if user_subtitle.audio else None
        webVit_url = fs.url(user_subtitle.webVit.name) if user_subtitle.webVit else None

        context = { 
            'subtitle_url': subtitle_url,
            'audio_url': audio_url,
            'webVit_url': webVit_url,
            'user_subtitle': user_subtitle
        }
        return render(request, 'view_subtitle.html', context)
    except Subtitle.DoesNotExist:
        messages.error(request, "Subtitle not found")
        return redirect('dashboard')

@login_required(login_url='/login')  
def save_transcript(request, pk):
    """Save edited transcript"""
    if request.method == 'POST':
        try:
            user_subtitle = Subtitle.objects.get(id=pk)
            
            # Check if user has permission to edit this subtitle
            if user_subtitle.user_profile.user != request.user:
                return JsonResponse({'status': 'error', 'message': 'Permission denied'}, status=403)
                
            data = json.loads(request.body)
            new_transcript = data.get('transcript')
            subtitle_offset = data.get('offset', 0)  # Get offset adjustment if provided

            if not new_transcript:
                return JsonResponse({'status': 'error', 'message': 'Transcript is empty.'}, status=400)
            
            # Get the file paths
            subtitle_path = user_subtitle.subtitle.path if user_subtitle.subtitle else None
            webvit_path = user_subtitle.webVit.path if user_subtitle.webVit else None
            
            if not subtitle_path or not webvit_path:
                return JsonResponse({'status': 'error', 'message': 'Subtitle files not found.'}, status=404)
            
            # Determine if the transcript is in SRT format (has timestamps)
            is_srt_format = ' --> ' in new_transcript
            
            # Update the SRT file
            if is_srt_format:
                # If offset adjustment is provided, apply it to all timestamps
                if subtitle_offset != 0:
                    try:
                        offset_ms = int(float(subtitle_offset) * 1000)  # Convert to milliseconds
                        
                        # Parse the SRT content and adjust timestamps
                        lines = new_transcript.split('\n')
                        adjusted_lines = []
                        
                        i = 0
                        while i < len(lines):
                            line = lines[i].strip()
                            adjusted_lines.append(line)
                            
                            # Check if this is a timestamp line
                            if '-->' in line:
                                # Parse start and end times
                                start_time, end_time = line.split(' --> ')
                                
                                # Convert to milliseconds, apply offset, and convert back
                                start_ms = srt_time_to_ms(start_time) + offset_ms
                                end_ms = srt_time_to_ms(end_time) + offset_ms
                                
                                # Ensure times don't go negative
                                start_ms = max(0, start_ms)
                                end_ms = max(start_ms + 500, end_ms)  # Ensure at least 500ms duration
                                
                                # Format back to SRT time
                                adjusted_start = ms_to_srt_time(start_ms)
                                adjusted_end = ms_to_srt_time(end_ms)
                                
                                # Replace the line with adjusted timestamps
                                adjusted_lines[-1] = f"{adjusted_start} --> {adjusted_end}"
                            
                            i += 1
                        
                        # Join lines back together
                        new_transcript = '\n'.join(adjusted_lines)
                    except Exception as e:
                        print(f"Error adjusting timestamps: {e}")
                        # Continue with original transcript if there's an error
                
                # If it's already in SRT format, write it directly
                with open(subtitle_path, 'w', encoding='utf-8') as f:
                    f.write(new_transcript)
                
                # Convert SRT to WebVTT for the webVit file
                vtt_content = "WEBVTT\n\n" + new_transcript.replace(',', '.')
                with open(webvit_path, 'w', encoding='utf-8') as f:
                    f.write(vtt_content)
            else:
                # If it's plain text, we need to convert it to SRT format
                # This is more complex and would require parsing the original SRT to maintain timestamps
                return JsonResponse({
                    'status': 'error', 
                    'message': 'Cannot save plain text as SRT. Please edit in SRT format with timestamps.'
                }, status=400)
            
            # Update the last modified time
            user_subtitle.updated_at = datetime.now()
            user_subtitle.save()
            
            return JsonResponse({
                'status': 'success', 
                'message': 'Subtitle files updated successfully!'
            })
            
        except Subtitle.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Subtitle not found.'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Invalid JSON data.'}, status=400)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error: {str(e)}'}, status=500)

    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'}, status=405)

@login_required(login_url='/login')
def adjust_subtitle_timing(request, pk):
    """Adjust subtitle timing with an offset"""
    if request.method == 'POST':
        try:
            user_subtitle = Subtitle.objects.get(id=pk)
            
            # Check if user has permission to edit this subtitle
            if user_subtitle.user_profile.user != request.user:
                return JsonResponse({'status': 'error', 'message': 'Permission denied'}, status=403)
            
            # Get the offset value in seconds
            offset = request.POST.get('offset', '0')
            try:
                offset_ms = int(float(offset) * 1000)  # Convert to milliseconds
            except ValueError:
                return JsonResponse({'status': 'error', 'message': 'Invalid offset value'}, status=400)
            
            # Get the file paths
            subtitle_path = user_subtitle.subtitle.path if user_subtitle.subtitle else None
            webvit_path = user_subtitle.webVit.path if user_subtitle.webVit else None
            
            if not subtitle_path or not webvit_path:
                return JsonResponse({'status': 'error', 'message': 'Subtitle files not found.'}, status=404)
            
            # Read the SRT file
            with open(subtitle_path, 'r', encoding='utf-8') as f:
                srt_content = f.read()
            
            # Parse the SRT content and adjust timestamps
            lines = srt_content.split('\n')
            adjusted_lines = []
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                adjusted_lines.append(line)
                
                # Check if this is a timestamp line
                if '-->' in line:
                    # Parse start and end times
                    start_time, end_time = line.split(' --> ')
                    
                    # Convert to milliseconds, apply offset, and convert back
                    start_ms = srt_time_to_ms(start_time) + offset_ms
                    end_ms = srt_time_to_ms(end_time) + offset_ms
                    
                    # Ensure times don't go negative
                    start_ms = max(0, start_ms)
                    end_ms = max(start_ms + 500, end_ms)  # Ensure at least 500ms duration
                    
                    # Format back to SRT time
                    adjusted_start = ms_to_srt_time(start_ms)
                    adjusted_end = ms_to_srt_time(end_ms)
                    
                    # Replace the line with adjusted timestamps
                    adjusted_lines[-1] = f"{adjusted_start} --> {adjusted_end}"
                
                i += 1
            
            # Join lines back together
            adjusted_srt = '\n'.join(adjusted_lines)
            
            # Write the adjusted SRT file
            with open(subtitle_path, 'w', encoding='utf-8') as f:
                f.write(adjusted_srt)
            
            # Convert SRT to WebVTT for the webVit file
            vtt_content = "WEBVTT\n\n" + adjusted_srt.replace(',', '.')
            with open(webvit_path, 'w', encoding='utf-8') as f:
                f.write(vtt_content)
            
            # Update the last modified time
            user_subtitle.updated_at = datetime.now()
            user_subtitle.save()
            
            return JsonResponse({
                'status': 'success', 
                'message': f'Subtitle timing adjusted by {offset} seconds'
            })
            
        except Subtitle.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Subtitle not found.'}, status=404)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error: {str(e)}'}, status=500)
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'}, status=405)

@login_required(login_url='/login')
def download_subtitle_format(request, pk, format_type):
    """Download subtitle in different formats"""
    try:
        user_subtitle = Subtitle.objects.get(id=pk)
        
        # Check if user has permission to download this subtitle
        if user_subtitle.user_profile.user != request.user:
            messages.error(request, "You don't have permission to download this subtitle")
            return redirect('dashboard')
            
        fs = FileSystemStorage()
        
        # Check if subtitle file exists
        if not user_subtitle.subtitle:
            return HttpResponse("No subtitle file found", status=404)
        
        # Get the subtitle file content
        subtitle_file = user_subtitle.subtitle.path
        with open(subtitle_file, 'r', encoding='utf-8') as f:
            srt_content = f.read()
        
        # Convert SRT to plain text
        plain_text = convert_srt_to_plain_text(srt_content)
        
        # Handle different format types
        if format_type == 'txt':
            # Return plain text file
            response = HttpResponse(plain_text, content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(user_subtitle.subtitle.name).replace(".srt", ".txt")}"'
            return response
        
        elif format_type == 'docx':
            # Create a DOCX file
            doc = Document()
            doc.add_heading(f"Transcript: {os.path.basename(user_subtitle.subtitle.name)}", 0)
            paragraphs = plain_text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save to a temporary file
            temp_file = BytesIO()
            doc.save(temp_file)
            temp_file.seek(0)
            
            # Return as a download
            response = HttpResponse(temp_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(user_subtitle.subtitle.name).replace(".srt", ".docx")}"'
            return response
        
        elif format_type == 'pdf':
            # Create a PDF file
            buffer = BytesIO()
            p = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Add title
            p.setFont("Helvetica-Bold", 16)
            p.drawString(72, height - 72, f"Transcript: {os.path.basename(user_subtitle.subtitle.name)}")
            
            # Add content
            p.setFont("Helvetica", 12)
            y_position = height - 100
            paragraphs = plain_text.split('\n\n')
            
            for paragraph in paragraphs:
                if not paragraph.strip():
                    continue
                    
                # Basic text wrapping (simplified)
                lines = []
                words = paragraph.split()
                current_line = ""
                
                for word in words:
                    if len(current_line + " " + word) * 6 < width - 144:  # Rough estimation of line width
                        current_line += " " + word if current_line else word
                    else:
                        lines.append(current_line)
                        current_line = word
                        
                if current_line:
                    lines.append(current_line)
                    
                # Draw text and update position
                for line in lines:
                    p.drawString(72, y_position, line)
                    y_position -= 15
                    
                    # Add a new page if needed
                    if y_position < 72:
                        p.showPage()
                        p.setFont("Helvetica", 12)
                        y_position = height - 72
                
                # Add paragraph spacing
                y_position -= 10
            
            p.showPage()
            p.save()
            buffer.seek(0)
            
            # Return as a download
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(user_subtitle.subtitle.name).replace(".srt", ".pdf")}"'
            return response
        
        elif format_type == 'srt':
            # Return original SRT file
            response = HttpResponse(srt_content, content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(user_subtitle.subtitle.name)}"'
            return response
            
        elif format_type == 'vtt':
            # Convert SRT to VTT
            vtt_content = "WEBVTT\n\n" + srt_content.replace(',', '.')
            response = HttpResponse(vtt_content, content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(user_subtitle.subtitle.name).replace(".srt", ".vtt")}"'
            return response
        
        else:
            return HttpResponse("Invalid format requested", status=400)
            
    except Subtitle.DoesNotExist:
        return HttpResponse("Subtitle not found", status=404)
    except Exception as e:
        return HttpResponse(f"Error: {str(e)}", status=500)