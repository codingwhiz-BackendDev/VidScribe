import tempfile
from django.shortcuts import render, redirect
from moviepy import VideoFileClip
from django.core.files.storage import FileSystemStorage
import os
import assemblyai as aai
import requests
import json
import time

# Set your AssemblyAI API key
aai.settings.api_key = "f501bc5824fc4140a87d4ba017621822"

# Using MyMemory Translation API - completely free for up to 5000 words/day
# No API key required for basic usage
MYMEMORY_API = "https://api.mymemory.translated.net/get"

def index(request):
    return render(request, 'index.html')

def ms_to_srt_time(ms):
    """Convert milliseconds to SRT time format (HH:MM:SS,mmm)"""
    seconds, ms = divmod(ms, 1000)
    minutes, seconds = divmod(seconds, 60)
    hours, minutes = divmod(minutes, 60)
    return f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d},{int(ms):03d}"

def translate_text(text, source_lang, target_lang):
    """Translate text using MyMemory Translation API"""
    # Function unchanged - keeping existing implementation
    if source_lang == target_lang:
        return text
    
    language_map = {
        'en': 'en',
        'hi': 'hi',
        'es': 'es',
        'fr': 'fr',
        'de': 'de',
        'ja': 'ja',
        'ru': 'ru',
        'zh': 'zh',
    }
    
    source = language_map.get(source_lang, source_lang)
    target = language_map.get(target_lang, target_lang)
    
    max_chunk_size = 500
    if len(text) > max_chunk_size:
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < max_chunk_size:
                current_chunk += sentence + ". "
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        translated_chunks = []
        for chunk in chunks:
            translated_chunk = translate_chunk(chunk, source, target)
            translated_chunks.append(translated_chunk)
            time.sleep(0.5)
        
        return " ".join(translated_chunks)
    else:
        return translate_chunk(text, source, target)

def translate_chunk(text, source_lang, target_lang):
    """Helper function to translate a single chunk of text"""
    params = {
        "q": text,
        "langpair": f"{source_lang}|{target_lang}",
        "de": "youremail@example.com"
    }
    
    try:
        response = requests.get(MYMEMORY_API, params=params)
        if response.status_code == 200:
            data = response.json()
            if data and 'responseData' in data and 'translatedText' in data['responseData']:
                return data['responseData']['translatedText']
        
        print(f"Translation error: {response.text}")
        return text + " [Translation failed]"
    except Exception as e:
        print(f"Translation exception: {str(e)}")
        return text + " [Translation error]"
def upload(request):
    if request.method == "POST":
        video = request.FILES.get('video')
        source_lang = request.POST.get('source-language')
        target_lang = request.POST.get('target-language')
        output_format = request.POST.get('format', 'srt')  # Default to srt if not specified
        include_timestamps = request.POST.get('timestamps') == 'on'  # Checkbox value
        transcription_quality = request.POST.get('quality', 'premium')  # Quality option
        speaker_id = request.POST.get('speaker-id') == 'on'  # Speaker identification
        
        fs = FileSystemStorage()
        
        # Process video in a temporary file
        with tempfile.NamedTemporaryFile(suffix=os.path.splitext(video.name)[1], delete=False) as temp_video:
            # Save uploaded video to temp file
            for chunk in video.chunks():
                temp_video.write(chunk)
            
            # Close the file so it can be accessed by moviepy
            temp_video.close()
            
            video_temp_path = temp_video.name
            base_name = os.path.splitext(os.path.basename(video.name))[0]
            
            # Extract audio and save it permanently
            audio_filename = f"{base_name}.mp3"
            audio_path = os.path.join(fs.location, audio_filename)
            
            # Load video and extract audio
            clip = VideoFileClip(video_temp_path).with_volume_scaled(1.5)
            clip.audio.write_audiofile(audio_path)
            
            # Free up resources
            clip.close()
            
            # Create a transcriber object
            transcriber = aai.Transcriber()
            
            # Set transcription config based on user choices
            config_params = {
                'language_code': source_lang,
                'punctuate': True,
                'format_text': True,
            }
            
            # Add speaker labels if requested and language is supported
            if speaker_id and source_lang in ['en']:
                config_params['speaker_labels'] = True
            
            # Adjust word boost based on quality
            if transcription_quality == 'premium' or transcription_quality == 'ultra':
                config_params['word_boost'] = ["div", "center", "CSS", "HTML"]
            
            # Request transcription
            transcript = transcriber.transcribe(
                audio_path,
                config=aai.TranscriptionConfig(**config_params)
            )
            
            # Group words into subtitle chunks
            max_words_per_chunk = 7
            translated_subs = []
            
            if hasattr(transcript, 'words') and transcript.words:
                current_chunk = []
                chunks = []
                
                for word in transcript.words:
                    current_chunk.append(word)
                    
                    if (len(current_chunk) >= max_words_per_chunk or 
                        (word.text in ['.', '!', '?', '...'] and len(current_chunk) > 1)):
                        chunks.append(current_chunk)
                        current_chunk = []
                
                if current_chunk:
                    chunks.append(current_chunk)
                
                # Process and translate each chunk
                for chunk in chunks:
                    if not chunk:
                        continue
                        
                    chunk_text = ' '.join(word.text for word in chunk)
                    chunk_text = chunk_text.replace(' .', '.').replace(' ,', ',').replace(' ?', '?').replace(' !', '!')
                    
                    start_time = chunk[0].start
                    end_time = chunk[-1].end
                    
                    translated_text = translate_text(chunk_text, source_lang, target_lang)
                    
                    # Add speaker info if available and requested
                    if speaker_id and hasattr(chunk[0], 'speaker') and chunk[0].speaker:
                        translated_text = f"Speaker {chunk[0].speaker}: {translated_text}"
                    
                    translated_subs.append({
                        'start': start_time,
                        'end': end_time,
                        'text': translated_text
                    })
            else:
                # Fallback to utterances
                for utterance in transcript.utterances:
                    words = utterance.text.split()
                    for i in range(0, len(words), max_words_per_chunk):
                        chunk_words = words[i:i+max_words_per_chunk]
                        chunk_text = ' '.join(chunk_words)
                        
                        utterance_duration = utterance.end - utterance.start
                        words_fraction = len(chunk_words) / len(words)
                        chunk_duration = utterance_duration * words_fraction
                        
                        chunk_start = utterance.start + (i / len(words)) * utterance_duration
                        chunk_end = chunk_start + chunk_duration
                        
                        translated_text = translate_text(chunk_text, source_lang, target_lang)
                        
                        # Add speaker info if available and requested
                        if speaker_id and hasattr(utterance, 'speaker') and utterance.speaker:
                            translated_text = f"Speaker {utterance.speaker}: {translated_text}"
                        
                        translated_subs.append({
                            'start': chunk_start,
                            'end': chunk_end,
                            'text': translated_text
                        })
            
            # Generate output based on selected format
            output_filename = f"{base_name}_{target_lang}"
            output_path = None
            
            if output_format == 'txt':
                # Plain text output
                output_filename += '.txt'
                output_path = os.path.join(fs.location, output_filename)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    for sub in translated_subs:
                        if include_timestamps:
                            timestamp = f"[{ms_to_srt_time(sub['start']).replace(',', '.')} - {ms_to_srt_time(sub['end']).replace(',', '.')}] "
                            f.write(f"{timestamp}{sub['text']}\n\n")
                        else:
                            f.write(f"{sub['text']}\n\n")
            
            elif output_format == 'srt':
                # SRT subtitle format
                output_filename += '.srt'
                output_path = os.path.join(fs.location, output_filename)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    for i, sub in enumerate(translated_subs, start=1):
                        f.write(f"{i}\n")
                        f.write(f"{ms_to_srt_time(sub['start'])} --> {ms_to_srt_time(sub['end'])}\n")
                        f.write(f"{sub['text']}\n\n")
            
            elif output_format == 'vtt':
                # VTT subtitle format
                output_filename += '.vtt'
                output_path = os.path.join(fs.location, output_filename)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write("WEBVTT\n\n")
                    for i, sub in enumerate(translated_subs, start=1):
                        # VTT uses . instead of , for milliseconds
                        start = ms_to_srt_time(sub['start']).replace(',', '.')
                        end = ms_to_srt_time(sub['end']).replace(',', '.')
                        f.write(f"{start} --> {end}\n")
                        f.write(f"{sub['text']}\n\n")
            
            elif output_format == 'docx':
                # Word document format
                try:
                    from docx import Document
                    
                    output_filename += '.docx'
                    output_path = os.path.join(fs.location, output_filename)
                    
                    doc = Document()
                    doc.add_heading(f"Transcript: {base_name}", 0)
                    
                    for sub in translated_subs:
                        paragraph = doc.add_paragraph()
                        if include_timestamps:
                            timestamp = f"[{ms_to_srt_time(sub['start']).replace(',', '.')} - {ms_to_srt_time(sub['end']).replace(',', '.')}] "
                            paragraph.add_run(timestamp).bold = True
                        paragraph.add_run(sub['text'])
                    
                    doc.save(output_path)
                    
                except ImportError:
                    # Fallback to text if python-docx not installed
                    output_filename = f"{base_name}_{target_lang}.txt"
                    output_path = os.path.join(fs.location, output_filename)
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(f"Transcript: {base_name}\n\n")
                        for sub in translated_subs:
                            if include_timestamps:
                                timestamp = f"[{ms_to_srt_time(sub['start']).replace(',', '.')} - {ms_to_srt_time(sub['end']).replace(',', '.')}] "
                                f.write(f"{timestamp}{sub['text']}\n\n")
                            else:
                                f.write(f"{sub['text']}\n\n")
            
            elif output_format == 'pdf':
                # PDF format
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib import colors
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    
                    output_filename += '.pdf'
                    output_path = os.path.join(fs.location, output_filename)
                    
                    doc = SimpleDocTemplate(output_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    title_style = styles['Heading1']
                    normal_style = styles['Normal']
                    timestamp_style = ParagraphStyle(
                        'TimeStamp', 
                        parent=normal_style,
                        textColor=colors.gray
                    )
                    
                    story.append(Paragraph(f"Transcript: {base_name}", title_style))
                    story.append(Spacer(1, 12))
                    
                    for sub in translated_subs:
                        if include_timestamps:
                            timestamp = f"[{ms_to_srt_time(sub['start']).replace(',', '.')} - {ms_to_srt_time(sub['end']).replace(',', '.')}]"
                            story.append(Paragraph(timestamp, timestamp_style))
                        
                        story.append(Paragraph(sub['text'], normal_style))
                        story.append(Spacer(1, 6))
                    
                    doc.build(story)
                    
                except ImportError:
                    # Fallback to text if reportlab not installed
                    output_filename = f"{base_name}_{target_lang}.txt"
                    output_path = os.path.join(fs.location, output_filename)
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(f"Transcript: {base_name}\n\n")
                        for sub in translated_subs:
                            if include_timestamps:
                                timestamp = f"[{ms_to_srt_time(sub['start']).replace(',', '.')} - {ms_to_srt_time(sub['end']).replace(',', '.')}] "
                                f.write(f"{timestamp}{sub['text']}\n\n")
                            else:
                                f.write(f"{sub['text']}\n\n")
            
            # For browser viewing, always create a VTT file
            vtt_filename = f"{base_name}_{target_lang}.vtt"
            vtt_path = os.path.join(fs.location, vtt_filename)
            
            # If we didn't already create a VTT file above
            if output_format != 'vtt':
                with open(vtt_path, 'w', encoding='utf-8') as f:
                    f.write("WEBVTT\n\n")
                    for i, sub in enumerate(translated_subs, start=1):
                        start = ms_to_srt_time(sub['start']).replace(',', '.')
                        end = ms_to_srt_time(sub['end']).replace(',', '.')
                        f.write(f"{start} --> {end}\n")
                        f.write(f"{sub['text']}\n\n")
            
            # For video streaming, create a temporary link that will work during this session
            temp_video_filename = f"temp_{base_name}_{int(time.time())}.mp4"
            temp_video_path = os.path.join(fs.location, temp_video_filename)
            
            # Create a temporary copy for browser viewing
            import shutil
            shutil.copy2(video_temp_path, temp_video_path)
            
            # Set a flag to delete the temp video after viewing
            request.session['temp_video_to_delete'] = temp_video_path
            
            # Clean up the initial temporary file
            try:
                os.unlink(video_temp_path)
            except Exception as e:
                print(f"Error removing initial temp file: {e}")
            
            # Pass data to the template
            context = {
                'video_url': fs.url(temp_video_filename),
                'subtitle_url': fs.url(vtt_filename),
                'output_url': fs.url(output_filename) if output_path else None,
                'output_format': output_format,
                'success': True
            }
            
            return render(request, 'upload.html', context)
    
    return render(request, 'upload.html')